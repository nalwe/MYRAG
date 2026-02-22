from django.db import transaction
from django.urls import reverse
from django.http import HttpResponse, HttpResponseForbidden
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required

from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document

import markdown

from rag.models import ChatSession, ChatMessage, ChatContext
from documents.models import Document as UserDocument
from documents.utils import get_accessible_documents
from rag.utils import create_onboarding_chat
from rag.retriever import retrieve_chunks
from rag.rag_pipeline import rag_answer
from accounts.models import OrganizationMember


# =====================================================
# 🔐 AI ACCESS CONTROL
# =====================================================

def check_ai_access(user):
    if not user.is_authenticated:
        return False, "Login required"

    if user.is_superuser:
        return True, None

    member = (
        OrganizationMember.objects
        .filter(user=user, is_active=True)
        .select_related("organization")
        .first()
    )

    if not member:
        return False, "AI access restricted"

    if not member.organization or not member.organization.is_active:
        return False, "Organization is inactive"

    return True, None


# =====================================================
# 💬 CHAT VIEW
# =====================================================

@login_required
def chat_view(request, session_id=None):
    user = request.user

    # 🔐 Access control
    allowed, error = check_ai_access(user)
    if not allowed:
        return HttpResponseForbidden(error)

    # 🧠 Ensure onboarding session exists
    create_onboarding_chat(user)

    sessions = ChatSession.objects.filter(user=user).order_by("-created_at")

    # 💬 Get or create session
    if session_id:
        session = get_object_or_404(ChatSession, id=session_id, user=user)
    else:
        session = sessions.first() or ChatSession.objects.create(user=user)

    messages = session.messages.order_by("created_at")

    # =====================================================
    # 📄 Active document (optional single-document scope)
    # =====================================================
    doc_id = request.GET.get("doc")
    active_document = None

    if doc_id:
        try:
            active_document = get_accessible_documents(user).get(id=doc_id)
        except Document.DoesNotExist:
            active_document = None

    # =====================================================
    # 📁 Folder-based restriction (session toggle)
    # =====================================================
    restrict_rag = request.session.get("restrict_rag", False)
    folder_id = request.GET.get("folder")
    active_folder = None

    if restrict_rag and folder_id:
        active_folder = Folder.objects.filter(
            id=folder_id,
            uploaded_by=user
        ).first()

    # =====================================================
    # ✉️ Handle message submission
    # =====================================================
    if request.method == "POST":
        query = request.POST.get("query", "").strip()

        if query:
            # 💾 Save user message
            ChatMessage.objects.create(
                session=session,
                role="user",
                content=query,
            )

            # ===============================
            # 🔎 Retrieval Logic
            # ===============================
            if active_document:
                # 📄 Single document scope
                retrieved = retrieve_chunks(
                    user=user,
                    query=query,
                    document_ids=[active_document.id],
                    k=5,
                )

            elif active_folder:
                # 📁 Folder-restricted scope
                folder_docs = get_accessible_documents(user).filter(
                    folder=active_folder
                )

                retrieved = retrieve_chunks(
                    user=user,
                    query=query,
                    document_ids=list(folder_docs.values_list("id", flat=True)),
                    k=5,
                )

            else:
                # 🌍 Global scope
                retrieved = retrieve_chunks(
                    user=user,
                    query=query,
                    k=5,
                )

            retrieved_texts = [r["text"] for r in retrieved]

            # 🤖 Generate grounded answer
            answer_md = rag_answer(
                question=query,
                chunks=retrieved,
            )

            answer_html = markdown.markdown(
                answer_md,
                extensions=["extra", "sane_lists"]
            )

            # 📚 Collect unique source documents
            source_docs = {}
            for r in retrieved:
                source_docs[r["document_id"]] = r["document_title"]

            formatted_sources = [
                {"id": doc_id, "title": title}
                for doc_id, title in source_docs.items()
            ]

            # 💾 Save assistant message
            ChatMessage.objects.create(
                session=session,
                role="assistant",
                content=answer_html,
                sources={
                    "documents": formatted_sources,
                    "chunks": retrieved_texts,
                    "restricted_to_folder": active_folder.id if active_folder else None,
                },
            )

        # 🔁 Preserve scope in redirect
        redirect_url = reverse("chat_session", args=[session.id])

        params = []
        if active_document:
            params.append(f"doc={active_document.id}")
        if active_folder:
            params.append(f"folder={active_folder.id}")

        if params:
            redirect_url += "?" + "&".join(params)

        return redirect(redirect_url)

    # =====================================================
    # 🖥️ Render chat
    # =====================================================
    return render(
        request,
        "chat/chat.html",
        {
            "sessions": sessions,
            "active_session": session,
            "messages": messages,
            "active_document": active_document,
            "active_folder": active_folder,
            "restrict_rag": restrict_rag,
        },
    )



# =====================================================
# ❌ DELETE CHAT
# =====================================================

@login_required
def delete_chat(request, session_id):
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    session.delete()
    return redirect("chat_view")


# =====================================================
# ➕ START CHAT WITH CONTEXT
# =====================================================

@login_required
@transaction.atomic
def start_chat_with_context(request):
    if request.method != "POST":
        return redirect("documents:document_list")

    document_ids = request.POST.getlist("documents")
    folder_ids = request.POST.getlist("folders")

    session = ChatSession.objects.create(user=request.user)

    documents = get_accessible_documents(request.user)

    if document_ids:
        documents = documents.filter(id__in=document_ids)

    if folder_ids:
        documents = documents.filter(folder_id__in=folder_ids)

    for doc in documents.distinct():
        ChatContext.objects.get_or_create(
            session=session,
            document=doc
        )

    return redirect("chat_view", session_id=session.id)


# =====================================================
# 📄 EXPORT CHAT PDF
# =====================================================

@login_required
def export_chat_pdf(request, session_id):
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    messages = session.messages.all()

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="chat_{session.id}.pdf"'

    pdf = canvas.Canvas(response)
    y = 800

    for msg in messages:
        pdf.drawString(40, y, f"{msg.role.upper()}:")
        y -= 20

        for line in msg.content.split("\n"):
            pdf.drawString(60, y, line[:120])
            y -= 15

        y -= 20
        if y < 100:
            pdf.showPage()
            y = 800

    pdf.save()
    return response


# =====================================================
# 📄 EXPORT CHAT DOCX
# =====================================================

@login_required
def export_chat_docx(request, session_id):
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    messages = session.messages.all()

    doc = Document()
    doc.add_heading(session.title or "Chat Export", level=1)

    for msg in messages:
        p = doc.add_paragraph()
        p.add_run(msg.role.upper()).bold = True
        doc.add_paragraph(msg.content)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="chat_{session.id}.docx"'
    doc.save(response)
    return response


# =====================================================
# 📄 EXPORT ANSWER PDF
# =====================================================

@login_required
def export_answer_pdf(request, message_id):
    message = get_object_or_404(
        ChatMessage,
        id=message_id,
        role="assistant",
        session__user=request.user
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    story = []

    soup = BeautifulSoup(message.content, "html.parser")

    for element in soup.find_all(["h2", "h3", "p", "li"]):
        story.append(Paragraph(element.get_text(), styles["Normal"]))

    doc.build(story)

    buffer.seek(0)
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=answer.pdf"

    return response


# =====================================================
# 📄 EXPORT ANSWER DOCX
# =====================================================

@login_required
def export_answer_docx(request, message_id):
    message = get_object_or_404(
        ChatMessage,
        id=message_id,
        role="assistant",
        session__user=request.user
    )

    doc = Document()
    soup = BeautifulSoup(message.content, "html.parser")

    for el in soup.find_all(["h2", "h3", "p", "li"]):
        doc.add_paragraph(el.get_text())

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=answer.docx"
    doc.save(response)
    return response


# =====================================================
# 🧹 CLEAR ALL CHATS
# =====================================================

@login_required
def clear_all_chats(request):
    if request.method == "POST":
        with transaction.atomic():
            ChatMessage.objects.filter(session__user=request.user).delete()
            ChatSession.objects.filter(user=request.user).delete()

    return redirect("chat_view")


# =====================================================
# 📄 EXPORT SELECTED MESSAGES PDF
# =====================================================

@login_required
def export_selected_messages_pdf(request, session_id):
    session = get_object_or_404(ChatSession, id=session_id, user=request.user)
    ids = request.POST.getlist("message_ids")
    messages = session.messages.filter(id__in=ids)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=selected_messages.pdf"

    pdf = canvas.Canvas(response)
    y = 800

    for msg in messages:
        pdf.drawString(40, y, f"{msg.role.upper()}:")
        y -= 18

        for line in msg.content.split("\n"):
            pdf.drawString(60, y, line[:120])
            y -= 14

        y -= 20
        if y < 100:
            pdf.showPage()
            y = 800

    pdf.save()
    return response
